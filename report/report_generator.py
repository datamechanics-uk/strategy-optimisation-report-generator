from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog, messagebox
import io
from .style_manager import StyleManager

class ReportGenerator:
    def __init__(self, intro_tab, param_tab, results_tab):
        self.intro_tab = intro_tab
        self.param_tab = param_tab
        self.results_tab = results_tab
        self.style_manager = None

    def generate(self):
        try:
            doc = Document("template.docx")
            self.style_manager = StyleManager(doc)
            self.style_manager.apply_styles()

            self.add_introduction_section(doc)
            self.add_parameter_sets_section(doc)
            self.add_results_section(doc)
            self.add_disclaimer_section(doc)

            self.save_document(doc)

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while generating the report: {str(e)}")

    def replace_placeholder(self, doc, placeholder, replacement):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)

    def add_introduction_section(self, doc):
        self.replace_placeholder(doc, "[strategy_name]", self.intro_tab.get_strategy_name())
        self.replace_placeholder(doc, "[target_performance_parameter]", self.intro_tab.get_target_performance())
        self.replace_placeholder(doc, "[optimisation_method]", self.intro_tab.get_optimisation_method())

        for paragraph in doc.paragraphs:
            if "[intro_table]" in paragraph.text:
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                
                headers = ["Market", "Timeframe", "Data Source", "Optimisation Timespan", "Out-of-Sample Timespan"]
                for i, header in enumerate(headers):
                    table.cell(0, i).text = header

                for row_data in self.intro_tab.get_intro_table_data():
                    cells = table.add_row().cells
                    for i, value in enumerate(row_data):
                        cells[i].text = str(value)

                p = paragraph._element
                p.getparent().replace(p, table._element)
                break
        
    def add_results_section(self, doc):
        for i, paragraph in enumerate(doc.paragraphs):
            if "[results_table]" in paragraph.text:
                results_data = self.results_tab.get_results_data()

                for index, (market, data) in enumerate(results_data.items(), start=1):
                    market_title = doc.add_paragraph(f"3.{index} {market}")
                    self.style_manager.apply_style_to_paragraph(market_title, "Heading 2")
                    paragraph._element.addnext(market_title._element)

                    if data["notes"]:
                        notes_paragraph = doc.add_paragraph(data["notes"])
                        market_title._element.addnext(notes_paragraph._element)
                        spacing_paragraph = doc.add_paragraph()
                        notes_paragraph._element.addnext(spacing_paragraph._element)
                    else:
                        spacing_paragraph = market_title

                    if data["equity_curve"]:
                        image_stream = io.BytesIO(data["equity_curve"])
                        img_paragraph = doc.add_paragraph()
                        img_paragraph.add_run().add_picture(image_stream, width=Inches(6))
                        spacing_paragraph._element.addnext(img_paragraph._element)
                        spacing_paragraph = img_paragraph

                    if data["performance_metrics"]:
                        image_stream = io.BytesIO(data["performance_metrics"])
                        img_paragraph = doc.add_paragraph()
                        img_paragraph.add_run().add_picture(image_stream, width=Inches(6))
                        spacing_paragraph._element.addnext(img_paragraph._element)

                    final_spacing_paragraph = doc.add_paragraph()
                    img_paragraph._element.addnext(final_spacing_paragraph._element)

                p = paragraph._element
                p.getparent().remove(p)
                break

    def add_disclaimer_section(self, doc):
        doc.add_page_break()

        disclaimer_title = doc.add_paragraph("Disclaimer")
        self.style_manager.apply_style_to_paragraph(disclaimer_title, "Heading 1")
        
        disclaimer_text = (
            "This document is provided for informational purposes only. Trading in the financial markets involves "
            "significant risk and may not be suitable for all investors. Past performance is not indicative of future results.\n\n"
            "Users of related software should conduct their own research and seek advice from a qualified financial advisor "
            "before making any investment decisions. This report and any associated software do not constitute financial "
            "advice and are not regulated by financial authorities.\n\n"
            "While efforts are made to ensure accuracy, the data, analysis, and any predictions or forecasts may not "
            "always be error-free and do not guarantee profits or predict market movements. Users are solely responsible "
            "for their investment decisions and any resulting outcomes.\n\n"
            "Data Mechanics Ltd. is not responsible for any financial losses incurred through the use of any associated "
            "software or any information included in this report or any related material. By using this report or "
            "associated software, you acknowledge and accept these terms and limitations."
        )   
        disclaimer_paragraph = doc.add_paragraph()
        run = disclaimer_paragraph.add_run(disclaimer_text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        disclaimer_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def create_market_title_style(self, doc):
        styles = doc.styles
        market_title_style = styles.add_style("MarketTitle", WD_STYLE_TYPE.PARAGRAPH)
        font = market_title_style.font
        font.name = "Arial"
        font.size = Pt(16)
        font.bold = False

    def add_parameter_sets_section(self, doc):
        for paragraph in doc.paragraphs:
            if "[parameter_set_table]" in paragraph.text:
                table = doc.add_table(rows=1, cols=7)
                table.style = "Table Grid"
                
                headers = ["Name", "Description", "Default", "Start", "Step", "End", "Best"]
                for i, header in enumerate(headers):
                    table.cell(0, i).text = header

                # The parameter data is expected to be a list of lists, where each inner list
                # contains 7 elements corresponding to the 7 columns of the table.
                # Each element represents: [Name, Description, Default, Start, Step, End, Best]
                for row_data in self.param_tab.get_parameter_data():
                    cells = table.add_row().cells
                    for i, value in enumerate(row_data):
                        cells[i].text = str(value)

                p = paragraph._element
                p.getparent().replace(p, table._element)
                break

    def save_document(self, doc):
        save_path = filedialog.asksaveasfilename(defaultextension=".docx")
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("Success", "Report generated successfully!")
        else:
            messagebox.showwarning("Cancelled", "Report generation cancelled.")